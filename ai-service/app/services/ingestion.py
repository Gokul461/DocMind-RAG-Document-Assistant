import asyncio
from fileinput import filename
from typing import AsyncGenerator
from uuid import uuid4

import pymupdf
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_huggingface import HuggingFaceEmbeddings

from app.core.config import get_settings

settings = get_settings()


def _get_embeddings() -> HuggingFaceEmbeddings:
    return HuggingFaceEmbeddings(
        model_name=settings.embedding_model,
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )


def _get_vectorstore() -> Chroma:
    return Chroma(
        collection_name="documents",
        embedding_function=_get_embeddings(),
        persist_directory=settings.chroma_persist_dir,
    )


def parse_pdf(file_bytes: bytes) -> list:
    doc = pymupdf.open(stream=file_bytes, filetype="pdf")
    pages = []
    for page_num, page in enumerate(doc, start=1):
        text = page.get_text("text").strip()
        if text:
            pages.append({"page": page_num, "text": text})
    doc.close()
    return pages


def parse_docx(file_bytes: bytes) -> list:
    import io
    doc = DocxDocument(io.BytesIO(file_bytes))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    pages = []
    page_size = 30
    for i in range(0, len(paragraphs), page_size):
        page_text = "\n".join(paragraphs[i : i + page_size])
        pages.append({"page": i // page_size + 1, "text": page_text})
    return pages


def parse_txt(file_bytes: bytes) -> list:
    text = file_bytes.decode("utf-8", errors="replace")
    pages = []
    size = 3000
    for i in range(0, len(text), size):
        pages.append({"page": i // size + 1, "text": text[i : i + size].strip()})
    return pages


PARSERS = {
    "pdf": parse_pdf,
    "docx": parse_docx,
    "txt": parse_txt,
}


def chunk_pages(pages: list, doc_id: str, filename: str = "") -> list:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        separators=["\n\n", "\n", ". ", "! ", "? ", " "],
    )
    documents = []
    for page in pages:
        splits = splitter.split_text(page["text"])
        for i, split in enumerate(splits):
            documents.append(
                Document(
                    page_content=split,
                    metadata={
                        "doc_id": doc_id,
                        "page": page["page"],
                        "chunk_index": i,
                        "filename": filename,
                    },
                )
            )
    return documents


async def ingest_document(
    file_bytes: bytes,
    filename: str,
) -> AsyncGenerator[dict, None]:
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext not in PARSERS:
        yield {"status": "error", "message": f"Unsupported file type: .{ext}"}
        return

    doc_id = str(uuid4())

    yield {"status": "parsing", "progress": 10, "message": f"Parsing {filename}...", "doc_id": doc_id}
    pages = await asyncio.to_thread(PARSERS[ext], file_bytes)
    if not pages:
        yield {"status": "error", "message": "Could not extract any text from file."}
        return

    yield {"status": "chunking", "progress": 35, "message": f"Splitting {len(pages)} pages into chunks...", "doc_id": doc_id}
    documents = await asyncio.to_thread(chunk_pages, pages, doc_id, filename)

    yield {"status": "embedding", "progress": 60, "message": f"Embedding {len(documents)} chunks locally...", "doc_id": doc_id}

    vectorstore = _get_vectorstore()
    await asyncio.to_thread(vectorstore.add_documents, documents)

    yield {
        "status": "complete",
        "progress": 100,
        "message": "Ready - you can now ask questions about this document.",
        "doc_id": doc_id,
        "filename": filename,
        "stats": {"pages": len(pages), "chunks": len(documents)},
    }
